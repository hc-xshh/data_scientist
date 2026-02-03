from langchain.tools import tool
import pandas as pd
import json
from pathlib import Path
from typing import Dict, Any, List
import re

# PDF处理
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word处理
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def pdf_to_markdown(file_path: str) -> str:
    """将PDF转换为Markdown格式"""
    if not PDF_AVAILABLE:
        return "错误：需要安装pdfplumber: pip install pdfplumber"
    
    try:
        markdown_content = []
        markdown_content.append(f"# PDF文档: {Path(file_path).name}\n")
        
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            markdown_content.append(f"**总页数**: {total_pages}\n")
            
            for i, page in enumerate(pdf.pages, 1):
                markdown_content.append(f"\n---\n## 第 {i} 页\n")
                
                # 提取文本
                text = page.extract_text()
                if text:
                    # 清理文本
                    text = text.strip()
                    markdown_content.append(text + "\n")
                
                # 提取表格并转为Markdown表格
                tables = page.extract_tables()
                if tables:
                    for j, table in enumerate(tables, 1):
                        if table and len(table) > 0:
                            markdown_content.append(f"\n### 表格 {j}\n")
                            markdown_content.append(table_to_markdown(table))
        
        return "\n".join(markdown_content)
    
    except Exception as e:
        return f"PDF转换失败: {str(e)}"


def word_to_markdown(file_path: str) -> str:
    """将Word文档转换为Markdown格式"""
    if not DOCX_AVAILABLE:
        return "错误：需要安装python-docx: pip install python-docx"
    
    try:
        doc = Document(file_path)
        markdown_content = []
        markdown_content.append(f"# Word文档: {Path(file_path).name}\n")
        
        # 处理段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 根据样式转换标题
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                if level.isdigit():
                    markdown_content.append(f"{'#' * int(level)} {text}\n")
                else:
                    markdown_content.append(f"## {text}\n")
            else:
                markdown_content.append(f"{text}\n")
        
        # 处理表格
        if doc.tables:
            markdown_content.append("\n---\n## 文档中的表格\n")
            
            for i, table in enumerate(doc.tables, 1):
                markdown_content.append(f"\n### 表格 {i}\n")
                
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                
                if rows:
                    markdown_content.append(table_to_markdown(rows))
        
        return "\n".join(markdown_content)
    
    except Exception as e:
        return f"Word转换失败: {str(e)}"


def csv_excel_to_markdown(file_path: str) -> str:
    """将CSV/Excel转换为Markdown格式"""
    try:
        ext = Path(file_path).suffix.lower()
        
        # 读取文件（读取全部数据）
        if ext == '.csv':
            # 尝试多种分隔符和编码
            df = None
            for encoding in ['utf-8', 'gbk', 'gb2312']:
                for sep in [',', '\t', ';', '|']:
                    try:
                        df = pd.read_csv(file_path, sep=sep, encoding=encoding)
                        if len(df.columns) > 1:
                            break
                    except:
                        continue
                if df is not None and len(df.columns) > 1:
                    break
            
            if df is None:
                df = pd.read_csv(file_path)  # 使用默认设置
                
        elif ext in ['.xlsx', '.xls']:
            df = pd.read_excel(file_path)
        else:
            return f"不支持的表格格式: {ext}"
        
        markdown_content = []
        markdown_content.append(f"# 数据文件: {Path(file_path).name}\n")
        
        # 数据摘要
        markdown_content.append("## 数据摘要\n")
        markdown_content.append(f"- **总行数**: {len(df):,}")
        markdown_content.append(f"- **总列数**: {len(df.columns)}")
        markdown_content.append(f"- **列名**: {', '.join(df.columns.tolist())}\n")
        
        # 数据质量
        missing_count = df.isnull().sum().sum()
        duplicate_count = df.duplicated().sum()
        if missing_count > 0 or duplicate_count > 0:
            markdown_content.append("### 数据质量")
            if missing_count > 0:
                markdown_content.append(f"- **缺失值**: {missing_count}")
                missing_cols = df.isnull().sum()
                missing_cols = missing_cols[missing_cols > 0]
                for col, count in missing_cols.items():
                    markdown_content.append(f"  - {col}: {count}")
            if duplicate_count > 0:
                markdown_content.append(f"- **重复行**: {duplicate_count}\n")
        
        # 数值列统计
        numeric_cols = df.select_dtypes(include='number').columns
        if len(numeric_cols) > 0:
            markdown_content.append("\n### 数值列统计\n")
            stats_df = df[numeric_cols].describe()
            markdown_content.append(dataframe_to_markdown(stats_df.round(2)))
        
        # 完整数据表格
        markdown_content.append("\n## 完整数据\n")
        markdown_content.append(dataframe_to_markdown(df))
        
        return "\n".join(markdown_content)
    
    except Exception as e:
        return f"表格转换失败: {str(e)}"


def json_to_markdown(file_path: str) -> str:
    """将JSON转换为Markdown格式"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        markdown_content = []
        markdown_content.append(f"# JSON文件: {Path(file_path).name}\n")
        
        # 如果是列表，转为表格
        if isinstance(data, list):
            markdown_content.append(f"**类型**: JSON数组\n")
            markdown_content.append(f"**元素数量**: {len(data)}\n")
            
            if len(data) > 0 and isinstance(data[0], dict):
                # 转为DataFrame再转为Markdown表格
                df = pd.DataFrame(data)
                markdown_content.append("## 数据内容\n")
                markdown_content.append(dataframe_to_markdown(df))
            else:
                # 直接展示
                markdown_content.append("## 数据内容\n")
                markdown_content.append("```json")
                markdown_content.append(json.dumps(data, ensure_ascii=False, indent=2))
                markdown_content.append("```")
        
        # 如果是字典
        elif isinstance(data, dict):
            markdown_content.append(f"**类型**: JSON对象\n")
            markdown_content.append(f"**顶层键**: {', '.join(data.keys())}\n")
            markdown_content.append("## 数据内容\n")
            markdown_content.append("```json")
            markdown_content.append(json.dumps(data, ensure_ascii=False, indent=2))
            markdown_content.append("```")
        
        else:
            markdown_content.append("## 数据内容\n")
            markdown_content.append(f"```\n{str(data)}\n```")
        
        return "\n".join(markdown_content)
    
    except Exception as e:
        return f"JSON转换失败: {str(e)}"


def text_to_markdown(file_path: str) -> str:
    """将文本文件转换为Markdown格式"""
    try:
        content = None
        detected_encoding = None
        
        # 尝试多种编码读取全部内容
        for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    detected_encoding = encoding
                    break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            return "无法读取文件：编码识别失败"
        
        markdown_content = []
        markdown_content.append(f"# 文本文件: {Path(file_path).name}\n")
        markdown_content.append(f"**编码**: {detected_encoding}")
        markdown_content.append(f"**总行数**: {len(content.splitlines()):,}")
        markdown_content.append(f"**总字符数**: {len(content):,}\n")
        markdown_content.append("---\n")
        markdown_content.append("## 文件内容\n")
        
        # 如果是代码文件，添加代码块
        ext = Path(file_path).suffix.lower()
        code_extensions = {
            '.py': 'python', '.js': 'javascript', '.java': 'java',
            '.cpp': 'cpp', '.c': 'c', '.sql': 'sql', '.sh': 'bash',
            '.html': 'html', '.css': 'css', '.xml': 'xml'
        }
        
        if ext in code_extensions:
            markdown_content.append(f"```{code_extensions[ext]}")
            markdown_content.append(content)
            markdown_content.append("```")
        else:
            markdown_content.append(content)
        
        return "\n".join(markdown_content)
    
    except Exception as e:
        return f"文本转换失败: {str(e)}"


def table_to_markdown(table_data: List[List]) -> str:
    """将二维列表转换为Markdown表格"""
    if not table_data or len(table_data) == 0:
        return ""
    
    markdown_lines = []
    
    # 表头
    header = table_data[0]
    markdown_lines.append("| " + " | ".join(str(cell) for cell in header) + " |")
    markdown_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    
    # 数据行
    for row in table_data[1:]:
        # 确保行长度与表头一致
        row_data = list(row) + [''] * (len(header) - len(row))
        markdown_lines.append("| " + " | ".join(str(cell) for cell in row_data[:len(header)]) + " |")
    
    return "\n".join(markdown_lines) + "\n"


def dataframe_to_markdown(df: pd.DataFrame) -> str:
    """将DataFrame转换为Markdown表格"""
    # 重置索引使其成为列
    df_reset = df.reset_index()
    
    # 转换为列表
    table_data = [df_reset.columns.tolist()] + df_reset.values.tolist()
    
    return table_to_markdown(table_data)


@tool
def analyze_file(file_path: str) -> str:
    """
    通用文件分析工具，将文件转换为Markdown格式并读取全部内容。
    
    支持格式：
    - 表格: CSV, Excel (.xlsx, .xls) - 读取全部数据
    - 文档: PDF, Word (.docx) - 提取所有文本和表格
    - 数据: JSON - 完整内容
    - 文本: TXT, MD, 代码文件等 - 完整内容
    
    返回完整的Markdown格式内容，便于AI理解和处理。
    """
    try:
        # 处理URL下载
        if file_path.startswith(('http://', 'https://')):
            import requests
            import tempfile
            
            response = requests.get(file_path, timeout=30)
            response.raise_for_status()
            
            suffix = Path(file_path).suffix or '.tmp'
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(response.content)
                file_path = tmp.name
        
        # 检查文件是否存在
        if not Path(file_path).exists():
            return f"错误：文件不存在 - {file_path}"
        
        # 获取文件信息
        file_stat = Path(file_path).stat()
        ext = Path(file_path).suffix.lower()
        
        # 根据文件类型转换为Markdown
        if ext == '.pdf':
            markdown_content = pdf_to_markdown(file_path)
        elif ext in ['.docx', '.doc']:
            markdown_content = word_to_markdown(file_path)
        elif ext in ['.csv', '.xlsx', '.xls']:
            markdown_content = csv_excel_to_markdown(file_path)
        elif ext == '.json':
            markdown_content = json_to_markdown(file_path)
        else:
            # 尝试作为文本文件处理
            markdown_content = text_to_markdown(file_path)
        
        # 添加文件元信息
        file_info_header = f"""---
**文件路径**: `{file_path}`  
**文件大小**: {file_stat.st_size / 1024:.2f} KB  
**文件类型**: {ext.upper()}  
---

"""
        
        return file_info_header + markdown_content
        
    except Exception as e:
        return f"# 文件分析失败\n\n**错误**: {type(e).__name__}\n\n**详情**: {str(e)}\n\n**文件**: {file_path}"


# 使用示例
if __name__ == "__main__":
    # 测试不同格式
    print(analyze_file("./data/sample.csv"))
    print("\n" + "="*80 + "\n")
    print(analyze_file("./documents/report.pdf"))