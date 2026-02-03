import os
from datetime import datetime
from typing import Dict, List, Optional, Union, Any
from langchain_core.tools import tool
import json
import tempfile

# Word文档相关依赖
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    WORD_AVAILABLE = True
except ImportError:
    print("⚠️ python-docx 未安装，Word功能不可用")
    WORD_AVAILABLE = False

FILE_FOLDER = tempfile.gettempdir()


def create_word_document_structure(content: Union[str, Dict, List], config: Dict) -> Dict:
    """
    创建Word文档结构
    
    Args:
        content: 文档内容，可以是字符串、字典或列表
        config: 文档配置（标题、作者、样式等）
    
    Returns:
        结构化的文档数据字典
    """
    # 如果是字符串，转换为结构
    if isinstance(content, str):
        try:
            # 尝试解析为JSON
            content_data = json.loads(content)
        except:
            # 如果是纯文本，创建简单结构
            content_data = {
                "title": config.get('title', 'AI生成文档'),
                "sections": [
                    {
                        "title": "内容",
                        "content": content
                    }
                ]
            }
    elif isinstance(content, list):
        # 如果是列表，转换为段落列表
        content_data = {
            "title": config.get('title', 'AI生成文档'),
            "sections": [
                {
                    "title": "内容列表",
                    "content": content
                }
            ]
        }
    else:
        # 已经是字典结构
        content_data = content
    
    # 确保有标题
    if 'title' not in content_data:
        content_data['title'] = config.get('title', 'AI生成文档')
    
    # 添加元数据
    if 'metadata' not in content_data:
        content_data['metadata'] = {}
    
    content_data['metadata'].update({
        'author': config.get('author', 'AI Assistant'),
        'generated_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'version': config.get('version', '1.0'),
        'generator': 'AI Document Generator'
    })
    
    return content_data


def create_word_document(content_data: Dict, config: Dict) -> Document:
    """
    创建Word文档对象
    
    Args:
        content_data: 结构化文档数据
        config: 文档配置
    
    Returns:
        Word文档对象
    """
    if not WORD_AVAILABLE:
        raise ImportError("python-docx未安装")
    
    # 创建文档
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4宽度
    section.page_height = Cm(29.7)  # A4高度
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # 创建样式
    styles = doc.styles
    
    # 标题样式
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = '微软雅黑' if config.get('chinese_font', True) else 'Calibri'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(30)
    title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # 一级标题样式
    heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.name = '微软雅黑' if config.get('chinese_font', True) else 'Calibri'
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(0, 77, 153)  # 蓝色
    heading1_style.paragraph_format.space_before = Pt(20)
    heading1_style.paragraph_format.space_after = Pt(10)
    heading1_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # 二级标题样式
    heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.name = '微软雅黑' if config.get('chinese_font', True) else 'Calibri'
    heading2_style.font.size = Pt(16)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(51, 102, 153)  # 浅蓝色
    heading2_style.paragraph_format.space_before = Pt(15)
    heading2_style.paragraph_format.space_after = Pt(8)
    heading2_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # 正文样式
    body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = '宋体' if config.get('chinese_font', True) else 'Times New Roman'
    body_style.font.size = Pt(12)
    body_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    body_style.paragraph_format.space_after = Pt(6)
    body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    body_style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进
    
    # 代码样式
    code_style = styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    code_style.font.color.rgb = RGBColor(46, 139, 87)  # 绿色
    code_style.paragraph_format.left_indent = Cm(1)
    code_style.paragraph_format.space_before = Pt(5)
    code_style.paragraph_format.space_after = Pt(5)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # 添加文档标题
    title = content_data.get('title', 'AI生成文档')
    title_para = doc.add_paragraph(title, style='CustomTitle')
    
    # 添加元数据
    if 'metadata' in content_data:
        metadata = content_data['metadata']
        meta_text = []
        if 'author' in metadata:
            meta_text.append(f"作者: {metadata['author']}")
        if 'generated_date' in metadata:
            meta_text.append(f"生成时间: {metadata['generated_date']}")
        if 'version' in metadata:
            meta_text.append(f"版本: {metadata['version']}")
        
        if meta_text:
            meta_para = doc.add_paragraph(" | ".join(meta_text), style='CustomBody')
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加文档内容
    if 'sections' in content_data:
        for section in content_data['sections']:
            # 添加章节标题
            if 'title' in section:
                heading_para = doc.add_paragraph(section['title'], style='CustomHeading1')
            
            # 添加章节内容
            if 'content' in section:
                content = section['content']
                
                if isinstance(content, list):
                    for item in content:
                        if isinstance(item, dict):
                            # 处理不同类型的内容
                            item_type = item.get('type', 'text')
                            if item_type == 'text':
                                doc.add_paragraph(item.get('text', ''), style='CustomBody')
                            elif item_type == 'code':
                                doc.add_paragraph(item.get('code', ''), style='CustomCode')
                            elif item_type == 'table':
                                if 'data' in item:
                                    _add_word_table(doc, item['data'])
                            elif item_type == 'list':
                                if 'items' in item:
                                    _add_word_list(doc, item['items'], item.get('ordered', False))
                        else:
                            doc.add_paragraph(str(item), style='CustomBody')
                else:
                    # 纯文本内容
                    paragraphs = str(content).split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip(), style='CustomBody')
            
            # 章节间空行
            doc.add_paragraph()
    
    elif 'content' in content_data:
        # 直接内容
        content = content_data['content']
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item), style='CustomBody')
        else:
            paragraphs = str(content).split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip(), style='CustomBody')
    
    # 添加页脚
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"第 {1} 页"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def _add_word_table(doc: Document, table_data: List[List]):
    """
    添加Word表格
    
    Args:
        doc: Word文档对象
        table_data: 表格数据
    """
    if not table_data:
        return
    
    rows = len(table_data)
    cols = len(table_data[0])
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Shading Accent 1'  # 使用预定义样式
    
    # 填充数据
    for i in range(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = str(table_data[i][j])
            
            # 设置表头样式
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _add_word_list(doc: Document, items: List, ordered: bool = False):
    """
    添加Word列表
    
    Args:
        doc: Word文档对象
        items: 列表项
        ordered: 是否有序列表
    """
    for i, item in enumerate(items):
        if ordered:
            para = doc.add_paragraph(f"{i+1}. {item}", style='CustomBody')
        else:
            para = doc.add_paragraph(f"• {item}", style='CustomBody')
        para.paragraph_format.left_indent = Cm(0.74)


@tool
def generate_word_document(
    content: Union[str, Dict, List],
    title: str = "AI生成文档",
    author: str = "AI Assistant",
    template_type: str = "professional",
    output_filename: Optional[str] = None,
    save_dir: Optional[str] = None
) -> str:
    """
    生成Word格式的文档
    
    参数:
        content: 文档内容，可以是：
            - 文本字符串
            - JSON字符串
            - 字典结构
            - 列表
            
            字典结构示例:
            {
                "title": "文档标题",
                "metadata": {
                    "author": "作者",
                    "version": "1.0"
                },
                "sections": [
                    {
                        "title": "章节1",
                        "content": [
                            {"type": "text", "text": "段落内容"},
                            {"type": "list", "items": ["项1", "项2"]},
                            {"type": "code", "code": "print('Hello')"},
                            {"type": "table", "data": [["标题1", "标题2"], ["数据1", "数据2"]]}
                        ]
                    }
                ]
            }
        
        title: 文档标题
        author: 作者
        template_type: 模板类型 (professional/modern/simple)
        output_filename: 输出文件名（不指定则自动生成）
        save_dir: 保存目录（不指定则使用默认temp目录）
    
    返回:
        生成的Word文档文件路径和下载URL
    
    示例:
        generate_word_document(
            content="这是一个测试文档内容",
            title="测试文档",
            author="张三"
        )
    """
    try:
        if not WORD_AVAILABLE:
            return "❌ 错误: python-docx库未安装，无法生成Word文档。请安装: pip install python-docx"
        
        # 准备保存目录
        if save_dir is None:
            save_dir = FILE_FOLDER
        
        os.makedirs(save_dir, exist_ok=True)
        
        # 生成文件名
        if output_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title[:50]  # 限制长度
            output_filename = f"{safe_title}_{timestamp}.docx"
        
        if not output_filename.endswith('.docx'):
            output_filename += '.docx'
        
        output_path = os.path.join(save_dir, output_filename)
        
        # 创建文档配置
        config = {
            'title': title,
            'author': author,
            'template_type': template_type,
            'chinese_font': True
        }
        
        # 创建文档结构
        content_data = create_word_document_structure(content, config)
        
        # 生成Word文档
        doc = create_word_document(content_data, config)
        
        # 保存文档
        doc.save(output_path)
        
        print(f"✅ Word文档已生成: {output_path}")
        print(f"📄 文件大小: {os.path.getsize(output_path):,} 字节")
        print(f"📝 文档标题: {title}")
        print(f"👤 作者: {author}")
        
        # 生成下载URL
        file_url = f"http://localhost:5000/download/{os.path.basename(output_path)}"
        preview_url = f"http://localhost:5000/preview/{os.path.basename(output_path)}"
        
        return f"""
🎉 Word文档生成成功！

📋 文档信息：
• 文档标题：{title}
• 作    者：{author}
• 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
• 文件格式：DOCX (Word文档)
• 文件大小：{os.path.getsize(output_path):,} 字节

📎 文件访问：
• 📍 本地路径：{output_path}
• ⬇️  下载链接：{file_url}
• 👁️  在线预览：{preview_url}

💡 使用说明：
1. 点击下载链接可以直接保存文档
2. 使用预览链接可以在浏览器中查看
3. 文档已按专业模板格式化，可直接使用
4. 支持Microsoft Word和WPS Office打开

🔄 如需重新生成或转换为其他格式，请告诉我！
"""
    
    except Exception as e:
        error_msg = f"❌ 生成Word文档失败: {str(e)}"
        print(error_msg)
        return error_msg


@tool
def quick_word_generate(
    text: str,
    title: str = "快速生成文档",
    author: str = "AI Assistant"
) -> str:
    """
    快速生成Word文档（简化接口）
    
    参数:
        text: 文档内容文本
        title: 文档标题
        author: 作者信息
    
    返回:
        Word文档下载链接
    
    示例:
        quick_word_generate(
            text="这是一个快速生成的测试文档...",
            title="测试报告",
            author="李四"
        )
    """
    return generate_word_document.invoke({
        "content": text,
        "title": title,
        "author": author,
        "template_type": "simple"
    })


# 测试函数
def test_word_generation():
    """测试Word文档生成"""
    print("🧪 测试Word文档生成工具...")
    
    # 测试1：简单文本
    print("\n📝 测试1：简单文本生成")
    result = generate_word_document.invoke({
        "content": "这是一个简单的测试文档。\n包含多行内容。\n第一行。\n第二行。",
        "title": "测试文档",
        "author": "测试用户"
    })
    print(f"结果: {result}")
    
    # 测试2：结构化内容
    print("\n📊 测试2：结构化内容生成")
    structured_content = {
        "title": "项目报告",
        "metadata": {
            "author": "项目组",
            "department": "技术部",
            "version": "1.0"
        },
        "sections": [
            {
                "title": "项目概述",
                "content": [
                    {"type": "text", "text": "这是一个AI文档生成项目。"},
                    {"type": "text", "text": "目标是开发智能文档生成系统。"},
                    {"type": "list", "items": [
                        "支持多种文档格式",
                        "提供专业模板",
                        "易于集成使用"
                    ]}
                ]
            },
            {
                "title": "技术实现",
                "content": [
                    {"type": "text", "text": "基于Python开发，使用python-docx库。"},
                    {"type": "code", "code": "def generate_document(content):\n    doc = Document()\n    doc.add_paragraph(content)\n    return doc"},
                    {"type": "table", "data": [
                        ["模块", "功能", "状态"],
                        ["文档生成", "生成Word文档", "已完成"],
                        ["格式设置", "设置样式布局", "进行中"],
                        ["模板支持", "支持多种模板", "规划中"]
                    ]}
                ]
            }
        ]
    }
    
    result = generate_word_document.invoke({
        "content": json.dumps(structured_content, ensure_ascii=False),
        "title": "结构化测试文档",
        "author": "开发团队",
        "template_type": "professional"
    })
    print(f"结果: {result}")


# if __name__ == "__main__":
#     test_word_generation()